#!/usr/bin/env python3
"""Convert V-model markdown documents to professional .docx files.

Requires: python-docx (pip install python-docx)
"""

import argparse
import json
import os
import re
import sys
from datetime import datetime, timezone

try:
	from docx import Document
	from docx.shared import Pt, RGBColor
	from docx.enum.text import WD_ALIGN_PARAGRAPH
	from docx.enum.table import WD_TABLE_ALIGNMENT
	from docx.oxml.ns import qn
	from docx.oxml import OxmlElement
except ImportError:
	print(
		"ERROR: python-docx is required but not installed.\n"
		"  Install it with: pip install python-docx",
		file=sys.stderr,
	)
	sys.exit(1)

MANIFEST_REL_PATH = os.path.join("docs", "v-model", "v-model-manifest.json")

DOC_TYPE_TITLES = {
	"vp": "Validation Plan",
	"urs": "User Requirements Specification",
	"fs": "Functional Specification",
	"ds": "Design Specification",
	"iq": "Installation Qualification Protocol",
	"oq": "Operational Qualification Protocol",
	"pq": "Performance Qualification Protocol",
	"traceability": "Traceability Matrix",
	"risk_assessment": "Risk Assessment",
	"vendor_assessment": "Vendor Assessment Package",
	"vsr": "Validation Summary Report — Vendor Evidence Summary",
}

def load_manifest(path: str) -> dict:
	with open(path, "r", encoding="utf-8") as fh:
		return json.load(fh)


def _find_md_file(doc_type: str, entry: dict, manifest_dir: str) -> str | None:
	"""Resolve a markdown file path from a manifest entry. Returns None if not found."""
	if entry.get("path"):
		candidate = os.path.normpath(os.path.join(manifest_dir, "..", "..", entry["path"]))
		if os.path.isfile(candidate):
			return os.path.abspath(candidate)
	candidate = os.path.join(manifest_dir, f"{doc_type}.md")
	if os.path.isfile(candidate):
		return os.path.abspath(candidate)
	return None


def resolve_input(input_arg: str, manifest: dict, manifest_dir: str) -> str:
	"""Resolve --input to an absolute markdown path (file path or doc-type key)."""
	if os.path.isfile(input_arg):
		return os.path.abspath(input_arg)
	doc_type = input_arg.lower()
	entry = manifest.get("documents", {}).get(doc_type, {})
	result = _find_md_file(doc_type, entry, manifest_dir)
	if result:
		return result
	print(f"ERROR: Cannot resolve input '{input_arg}' to a markdown file.", file=sys.stderr)
	sys.exit(1)


def doc_type_from_path(md_path: str) -> str:
	"""Infer doc-type key from a markdown filename."""
	stem = os.path.splitext(os.path.basename(md_path))[0].lower()
	for key in DOC_TYPE_TITLES:
		if stem == key or stem.startswith(key + "-") or stem.startswith(key + "_"):
			return key
	return stem


def _publishable_docs(manifest: dict, manifest_dir: str):
	"""Yield (doc_type, md_path) for documents with status != 'not_started'."""
	for doc_type, entry in manifest.get("documents", {}).items():
		if entry.get("status") == "not_started":
			continue
		md_path = _find_md_file(doc_type, entry, manifest_dir)
		if md_path:
			yield doc_type, md_path

_RE_HEADING = re.compile(r"^(#{1,4})\s+(.+)$")
_RE_TABLE_SEP = re.compile(r"^\|[\s\-:|]+\|$")
_RE_TABLE_ROW = re.compile(r"^\|(.+)\|$")
_RE_BULLET = re.compile(r"^-\s+(.+)$")
_RE_NUMBERED = re.compile(r"^\d+\.\s+(.+)$")
_RE_CODE_FENCE = re.compile(r"^```")
_RE_INLINE = re.compile(
	r"(?P<todo><!--\s*TODO:\s*(.+?)\s*-->)"
	r"|(?P<bold>\*\*(.+?)\*\*)"
	r"|(?P<code>`([^`]+)`)"
)
# Italic is processed in a second pass after bold is removed, to avoid
# false positives with bullet lists and adjacent bold markers.
_RE_ITALIC = re.compile(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)")


def _apply_inline(paragraph, text: str) -> None:
	"""Parse inline markdown (bold, italic, code, TODO) into paragraph runs.

	Two-pass approach: first pass handles TODO, bold, and code. Second pass
	handles italic on remaining plain-text runs, avoiding false positives with
	bullet-list asterisks and adjacent bold markers.
	"""
	# --- Pass 1: TODO, bold, code ---
	segments = []  # list of (text, style) tuples
	pos = 0
	for m in _RE_INLINE.finditer(text):
		if m.start() > pos:
			segments.append((text[pos:m.start()], "plain"))
		if m.group("todo"):
			segments.append((m.group(2), "todo"))
		elif m.group("bold"):
			segments.append((m.group(4), "bold"))
		elif m.group("code"):
			segments.append((m.group(6), "code"))
		pos = m.end()
	trailing = text[pos:]
	if trailing or pos == 0:
		segments.append((trailing, "plain"))

	# --- Pass 2: italic within plain segments ---
	for seg_text, style in segments:
		if style == "plain":
			ipos = 0
			for im in _RE_ITALIC.finditer(seg_text):
				if im.start() > ipos:
					paragraph.add_run(seg_text[ipos:im.start()])
				paragraph.add_run(im.group(1)).italic = True
				ipos = im.end()
			remainder = seg_text[ipos:]
			if remainder or ipos == 0:
				paragraph.add_run(remainder)
		elif style == "todo":
			run = paragraph.add_run(seg_text)
			run.font.highlight_color = 7  # WD_COLOR_INDEX.YELLOW
			run.bold = True
		elif style == "bold":
			paragraph.add_run(seg_text).bold = True
		elif style == "code":
			run = paragraph.add_run(seg_text)
			run.font.name = "Courier New"
			run.font.size = Pt(9)


def _flush_table(doc, rows: list[list[str]]) -> None:
	if not rows:
		return
	cols = len(rows[0])
	table = doc.add_table(rows=0, cols=cols)
	table.style = "Table Grid"
	table.alignment = WD_TABLE_ALIGNMENT.CENTER
	for i, cells in enumerate(rows):
		row = table.add_row()
		for j, cell_text in enumerate(cells[:cols]):
			para = row.cells[j].paragraphs[0]
			_apply_inline(para, cell_text.strip())
			if i == 0:
				for run in para.runs:
					run.bold = True


def _flush_code(doc, lines: list[str]) -> None:
	if not lines:
		return
	para = doc.add_paragraph()
	for cl in lines:
		run = para.add_run(cl + "\n")
		run.font.name, run.font.size = "Courier New", Pt(9)
		run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def render_markdown(doc, md_text: str) -> None:
	"""Parse markdown and append formatted content to the document."""
	table_rows, code_lines, in_code = [], [], False
	for line in md_text.split("\n"):
		# Code fence toggle
		if _RE_CODE_FENCE.match(line.strip()):
			if in_code:
				_flush_code(doc, code_lines)
				code_lines, in_code = [], False
			else:
				_flush_table(doc, table_rows)
				table_rows, in_code = [], True
			continue
		if in_code:
			code_lines.append(line)
			continue

		stripped = line.strip()

		if not stripped:
			_flush_table(doc, table_rows)
			table_rows = []
			continue

		# Heading
		hm = _RE_HEADING.match(stripped)
		if hm:
			_flush_table(doc, table_rows)
			table_rows = []
			doc.add_heading(hm.group(2), level=min(len(hm.group(1)), 4))
			continue

		if _RE_TABLE_SEP.match(stripped):
			continue

		trm = _RE_TABLE_ROW.match(stripped)
		if trm:
			table_rows.append([c.strip() for c in trm.group(1).split("|")])
			continue

		_flush_table(doc, table_rows)
		table_rows = []

		bm = _RE_BULLET.match(stripped)
		if bm:
			_apply_inline(doc.add_paragraph(style="List Bullet"), bm.group(1))
			continue
		nm = _RE_NUMBERED.match(stripped)
		if nm:
			_apply_inline(doc.add_paragraph(style="List Number"), nm.group(1))
			continue

		_apply_inline(doc.add_paragraph(), stripped)

	_flush_table(doc, table_rows)
	_flush_code(doc, code_lines)

def _styled_table(doc, rows: int, cols: int, headers: list[str]):
	"""Create a Table Grid table with bold header row. Returns the table."""
	table = doc.add_table(rows=rows, cols=cols)
	table.style = "Table Grid"
	table.alignment = WD_TABLE_ALIGNMENT.CENTER
	for j, h in enumerate(headers):
		table.rows[0].cells[j].paragraphs[0].add_run(h).bold = True
	return table


def _centered_text(doc, text: str, size: int, after: int = 12) -> None:
	para = doc.add_paragraph()
	para.alignment = WD_ALIGN_PARAGRAPH.CENTER
	run = para.add_run(text)
	run.bold = True
	run.font.size = Pt(size)
	para.space_after = Pt(after)


def _add_title_page(doc, manifest: dict, doc_type: str) -> None:
	project = manifest.get("project", {})
	doc_entry = manifest.get("documents", {}).get(doc_type, {})
	_centered_text(doc, project.get("name", "Untitled Project"), 28)
	_centered_text(doc, DOC_TYPE_TITLES.get(doc_type, doc_type.upper()), 22, after=24)

	id_prefix = project.get("id_prefix", "")
	doc_id = f"{id_prefix}{doc_type.upper()}" if id_prefix else doc_type.upper()
	last_updated = doc_entry.get("last_updated", datetime.now(timezone.utc).strftime("%Y-%m-%d"))
	meta = [
		("Document ID", doc_id),
		("Revision", doc_entry.get("version", "0.1")),
		("Date", last_updated[:10]),
		("Status", doc_entry.get("status", "draft").replace("_", " ").title()),
		("GAMP Category", project.get("gamp_category", "N/A")),
	]
	table = doc.add_table(rows=len(meta), cols=2)
	table.style = "Table Grid"
	table.alignment = WD_TABLE_ALIGNMENT.CENTER
	for i, (label, value) in enumerate(meta):
		table.rows[i].cells[0].paragraphs[0].add_run(label).bold = True
		table.rows[i].cells[1].paragraphs[0].add_run(value)
	doc.add_page_break()


def _add_revision_history(doc, manifest: dict, doc_type: str) -> None:
	doc.add_heading("Revision History", level=1)
	entry = manifest.get("documents", {}).get(doc_type, {})
	table = _styled_table(doc, 2, 4, ["Version", "Date", "Description", "Author"])
	status = entry.get("status", "draft").replace("_", " ").title()
	table.rows[1].cells[0].paragraphs[0].add_run(entry.get("version", "0.1"))
	table.rows[1].cells[1].paragraphs[0].add_run(entry.get("last_updated", "")[:10])
	table.rows[1].cells[2].paragraphs[0].add_run(f"Initial {status.lower()}")
	doc.add_page_break()


def _add_approval_block(doc) -> None:
	doc.add_heading("Approval Signatures", level=1)
	table = _styled_table(doc, 4, 4, ["Role", "Name", "Signature", "Date"])
	for i, role in enumerate(["Prepared by", "Reviewed by", "Approved by"], start=1):
		table.rows[i].cells[0].paragraphs[0].add_run(role)
	doc.add_page_break()


def _add_toc_placeholder(doc) -> None:
	doc.add_heading("Table of Contents", level=1)
	run = doc.add_paragraph().add_run()
	for fld_type in ("begin", "separate", "end"):
		el = OxmlElement("w:fldChar")
		el.set(qn("w:fldCharType"), fld_type)
		run._r.append(el)
		if fld_type == "begin":
			instr = OxmlElement("w:instrText")
			instr.set(qn("xml:space"), "preserve")
			instr.text = ' TOC \\o "1-4" \\h \\z \\u '
			run._r.append(instr)
	r = doc.add_paragraph().add_run(
		"[Right-click and select Update Field to populate this table of contents]")
	r.italic = True
	r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)


def _add_traceability_appendix(doc, manifest: dict, doc_type: str) -> None:
	reqs = manifest.get("requirements", [])
	dt = doc_type.upper()
	relevant = [r for r in reqs
		if r.get("source_document") == doc_type
		or r.get("id", "").upper().startswith(dt + "-")
		or any(dt in x.upper() for x in r.get("traces_to", []))
		or any(dt in x.upper() for x in r.get("tested_by", []))]
	if not relevant:
		return
	doc.add_heading("Appendix: Traceability", level=1)
	table = _styled_table(doc, 1, 4, ["Requirement ID", "Description", "Traces To", "Tested By"])
	for req in relevant:
		row = table.add_row()
		row.cells[0].paragraphs[0].add_run(req.get("id", ""))
		text = req.get("text", "")
		row.cells[1].paragraphs[0].add_run(text[:117] + "..." if len(text) > 120 else text)
		row.cells[2].paragraphs[0].add_run(", ".join(req.get("traces_to", [])))
		row.cells[3].paragraphs[0].add_run(", ".join(req.get("tested_by", [])))

# ---------------------------------------------------------------------------
# Core publish functions
# ---------------------------------------------------------------------------

def _new_doc(template_path: str | None):
	if template_path and os.path.isfile(template_path):
		return Document(template_path)
	return Document()


def publish_single(md_path: str, manifest: dict, output_path: str,
                   template_path: str | None = None) -> str:
	"""Convert a single markdown file to .docx. Returns the output path."""
	doc = _new_doc(template_path)
	doc_type = doc_type_from_path(md_path)

	with open(md_path, "r", encoding="utf-8") as fh:
		md_text = fh.read()

	_add_title_page(doc, manifest, doc_type)
	_add_revision_history(doc, manifest, doc_type)
	_add_approval_block(doc)
	_add_toc_placeholder(doc)
	doc.add_page_break()
	render_markdown(doc, md_text)
	_add_traceability_appendix(doc, manifest, doc_type)

	os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)
	doc.save(output_path)
	return os.path.abspath(output_path)


def publish_full(manifest: dict, manifest_path: str, template_path: str | None) -> list[str]:
	"""Publish all documents with status != 'not_started'."""
	manifest_dir = os.path.dirname(os.path.abspath(manifest_path))
	outputs = []
	for doc_type, md_path in _publishable_docs(manifest, manifest_dir):
		out = os.path.join(manifest_dir, os.path.splitext(os.path.basename(md_path))[0] + ".docx")
		result = publish_single(md_path, manifest, out, template_path)
		outputs.append(result)
		print(f"  Published: {doc_type} -> {result}")
	return outputs


def publish_bundle(manifest: dict, manifest_path: str, template_path: str | None) -> list[str]:
	"""Publish full package plus a combined master document."""
	outputs = publish_full(manifest, manifest_path, template_path)
	if len(outputs) < 2:
		return outputs

	manifest_dir = os.path.dirname(os.path.abspath(manifest_path))
	project_name = manifest.get("project", {}).get("name", "Project")
	master_path = os.path.join(manifest_dir, f"{project_name.lower().replace(' ', '_')}_bundle.docx")

	master = _new_doc(template_path)
	_centered_text(master, project_name, 28)
	_centered_text(master, "Validation Documentation Bundle", 20, after=24)
	p = master.add_paragraph()
	p.alignment = WD_ALIGN_PARAGRAPH.CENTER
	p.add_run(f"Generated: {datetime.now(timezone.utc).strftime('%Y-%m-%d')}")
	master.add_page_break()
	_add_toc_placeholder(master)
	master.add_page_break()

	for doc_type, md_path in _publishable_docs(manifest, manifest_dir):
		master.add_heading(DOC_TYPE_TITLES.get(doc_type, doc_type.upper()), level=1)
		with open(md_path, "r", encoding="utf-8") as fh:
			render_markdown(master, fh.read())
		master.add_page_break()

	master.save(master_path)
	outputs.append(os.path.abspath(master_path))
	print(f"  Bundle: {master_path}")
	return outputs

def _build_parser() -> argparse.ArgumentParser:
	parser = argparse.ArgumentParser(
		description="Convert V-model markdown documents to .docx with regulated-environment formatting.",
	)
	parser.add_argument("--input", help="Markdown file path or doc type name (e.g. 'urs')")
	parser.add_argument("--manifest", default=MANIFEST_REL_PATH,
	                    help=f"Path to v-model-manifest.json (default: {MANIFEST_REL_PATH})")
	parser.add_argument("--output", help="Output .docx path (default: input path with .docx ext)")
	parser.add_argument("--template", default=os.path.join("assets", "docx_template.docx"),
	                    help="Path to .docx template (default: assets/docx_template.docx)")
	parser.add_argument("--mode", choices=("single", "full", "bundle"), default="single",
	                    help="Publish mode: single doc, full package, or bundle (default: single)")
	return parser


def main() -> None:
	parser = _build_parser()
	args = parser.parse_args()

	if not os.path.isfile(args.manifest):
		print(f"ERROR: Manifest not found at {args.manifest}", file=sys.stderr)
		print("Run init_project.py first, or pass --manifest with the correct path.", file=sys.stderr)
		sys.exit(1)

	manifest = load_manifest(args.manifest)
	manifest_dir = os.path.dirname(os.path.abspath(args.manifest))
	template = args.template if os.path.isfile(args.template) else None

	if args.mode == "full":
		outputs = publish_full(manifest, args.manifest, template)
		print(f"\nPublished {len(outputs)} document(s).")
	elif args.mode == "bundle":
		outputs = publish_bundle(manifest, args.manifest, template)
		print(f"\nBundle complete: {len(outputs)} file(s).")
	else:
		if not args.input:
			print("ERROR: --input is required for single-document mode.", file=sys.stderr)
			sys.exit(1)
		md_path = resolve_input(args.input, manifest, manifest_dir)
		out_path = args.output or os.path.splitext(md_path)[0] + ".docx"
		print(f"Published: {publish_single(md_path, manifest, out_path, template)}")


if __name__ == "__main__":
	main()
