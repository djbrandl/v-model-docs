"""
Generate a professional .docx template for v-model-docs.

Standalone one-off script that creates a minimal Word template with:
- Professional heading styles (Heading 1-4)
- Normal text style
- Table style with header row and alternating shading
- Header with company/document placeholders
- Footer with page numbering and document ID
- Letter page size with 1-inch margins
"""

import sys
import os

try:
	from docx import Document
	from docx.shared import Pt, Inches, Cm, RGBColor, Emu
	from docx.enum.text import WD_ALIGN_PARAGRAPH
	from docx.enum.section import WD_ORIENT
	from docx.enum.table import WD_TABLE_ALIGNMENT
	from docx.oxml.ns import qn, nsdecls
	from docx.oxml import parse_xml
except ImportError:
	print(
		"ERROR: python-docx is required but not installed.\n"
		"  Install it with: pip install python-docx",
		file=sys.stderr,
	)
	sys.exit(1)


# --- Constants ---
BLUE_ACCENT = RGBColor(0x1F, 0x4E, 0x79)  # Dark professional blue
LIGHT_BLUE = "D6E4F0"                       # Light blue for table headers
ALT_ROW_GRAY = "F2F2F2"                     # Alternating row shading
FONT_NAME = "Calibri"
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT_DIR = os.path.join(os.path.dirname(SCRIPT_DIR), "assets")
OUTPUT_PATH = os.path.join(OUTPUT_DIR, "docx_template.docx")


def configure_page_layout(doc: Document) -> None:
	"""Set page size to Letter and margins to 1 inch on all sides."""
	section = doc.sections[0]
	section.page_width = Inches(8.5)
	section.page_height = Inches(11)
	section.orientation = WD_ORIENT.PORTRAIT
	section.top_margin = Inches(1)
	section.bottom_margin = Inches(1)
	section.left_margin = Inches(1)
	section.right_margin = Inches(1)
	# Header/footer distance from edge
	section.header_distance = Inches(0.5)
	section.footer_distance = Inches(0.5)


def set_run_font(run, name=FONT_NAME, size=None, bold=False, color=None):
	"""Apply font properties to a run."""
	run.font.name = name
	# Force Calibri for East Asian text as well
	run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
	if size is not None:
		run.font.size = Pt(size)
	run.font.bold = bold
	if color is not None:
		run.font.color.rgb = color


def configure_heading_style(style, size, color, space_before, space_after):
	"""Configure a heading style with font, size, color, and spacing."""
	font = style.font
	font.name = FONT_NAME
	font.size = Pt(size)
	font.bold = True
	font.color.rgb = color
	pf = style.paragraph_format
	pf.space_before = Pt(space_before)
	pf.space_after = Pt(space_after)
	pf.keep_with_next = True
	# Set the East Asian font via XML to ensure Calibri everywhere
	rpr = style.element.find(qn("w:rPr"))
	if rpr is not None:
		fonts = rpr.find(qn("w:rFonts"))
		if fonts is not None:
			fonts.set(qn("w:eastAsia"), FONT_NAME)


def configure_styles(doc: Document) -> None:
	"""Set up Heading 1-4, Normal, and other base styles."""
	styles = doc.styles

	# --- Normal ---
	normal = styles["Normal"]
	normal.font.name = FONT_NAME
	normal.font.size = Pt(11)
	normal.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
	normal.paragraph_format.space_after = Pt(6)
	normal.paragraph_format.line_spacing = 1.15
	# East Asian font
	rpr = normal.element.find(qn("w:rPr"))
	if rpr is not None:
		fonts = rpr.find(qn("w:rFonts"))
		if fonts is not None:
			fonts.set(qn("w:eastAsia"), FONT_NAME)

	# --- Headings ---
	heading_config = {
		"Heading 1": (20, BLUE_ACCENT, 24, 12),
		"Heading 2": (16, BLUE_ACCENT, 18, 8),
		"Heading 3": (13, BLUE_ACCENT, 14, 6),
		"Heading 4": (11, BLUE_ACCENT, 12, 4),
	}
	for name, (size, color, sp_before, sp_after) in heading_config.items():
		style = styles[name]
		configure_heading_style(style, size, color, sp_before, sp_after)


def build_header(doc: Document) -> None:
	"""Add header with company name (left) and document title (right)."""
	section = doc.sections[0]
	header = section.header
	header.is_linked_to_previous = False

	# Use a single paragraph with a right-aligned tab stop
	paragraph = header.paragraphs[0]
	paragraph.paragraph_format.space_after = Pt(0)
	paragraph.paragraph_format.space_before = Pt(0)

	# Add a bottom border to the header paragraph
	pPr = paragraph._element.get_or_add_pPr()
	pBdr = parse_xml(
		f'<w:pBdr {nsdecls("w")}>'
		'  <w:bottom w:val="single" w:sz="4" w:space="4" w:color="1F4E79"/>'
		"</w:pBdr>"
	)
	pPr.append(pBdr)

	# Set a right-aligned tab stop at the right margin (6.5 inches from left margin)
	tabs = parse_xml(
		f'<w:tabs {nsdecls("w")}>'
		'  <w:tab w:val="right" w:pos="9360"/>'
		"</w:tabs>"
	)
	pPr.append(tabs)

	# Left-aligned company name
	run_left = paragraph.add_run("[Company Name]")
	set_run_font(run_left, size=9, bold=True, color=BLUE_ACCENT)

	# Tab character to push next text right
	run_tab = paragraph.add_run("\t")
	run_tab.font.size = Pt(9)

	# Right-aligned document title
	run_right = paragraph.add_run("[Document Title]")
	set_run_font(run_right, size=9, color=BLUE_ACCENT)


def build_footer(doc: Document) -> None:
	"""Add footer with 'Page X of Y' (center) and document ID (right)."""
	section = doc.sections[0]
	footer = section.footer
	footer.is_linked_to_previous = False

	paragraph = footer.paragraphs[0]
	paragraph.paragraph_format.space_after = Pt(0)
	paragraph.paragraph_format.space_before = Pt(0)

	# Top border on the footer paragraph
	pPr = paragraph._element.get_or_add_pPr()
	pBdr = parse_xml(
		f'<w:pBdr {nsdecls("w")}>'
		'  <w:top w:val="single" w:sz="4" w:space="4" w:color="1F4E79"/>'
		"</w:pBdr>"
	)
	pPr.append(pBdr)

	# Tab stops: center at 3.25 inches, right at 6.5 inches
	tabs = parse_xml(
		f'<w:tabs {nsdecls("w")}>'
		'  <w:tab w:val="center" w:pos="4680"/>'
		'  <w:tab w:val="right" w:pos="9360"/>'
		"</w:tabs>"
	)
	pPr.append(tabs)

	# Tab to center position
	run_tab1 = paragraph.add_run("\t")
	run_tab1.font.size = Pt(8)

	# "Page " text
	run_page = paragraph.add_run("Page ")
	set_run_font(run_page, size=8, color=RGBColor(0x66, 0x66, 0x66))

	# PAGE field
	fld_page = parse_xml(
		f'<w:fldSimple {nsdecls("w")} w:instr=" PAGE \\* MERGEFORMAT ">'
		"  <w:r>"
		"    <w:rPr>"
		f'      <w:sz w:val="16"/>'
		"    </w:rPr>"
		"    <w:t>1</w:t>"
		"  </w:r>"
		"</w:fldSimple>"
	)
	paragraph._element.append(fld_page)

	# " of " text
	run_of = paragraph.add_run(" of ")
	set_run_font(run_of, size=8, color=RGBColor(0x66, 0x66, 0x66))

	# NUMPAGES field
	fld_total = parse_xml(
		f'<w:fldSimple {nsdecls("w")} w:instr=" NUMPAGES \\* MERGEFORMAT ">'
		"  <w:r>"
		"    <w:rPr>"
		f'      <w:sz w:val="16"/>'
		"    </w:rPr>"
		"    <w:t>1</w:t>"
		"  </w:r>"
		"</w:fldSimple>"
	)
	paragraph._element.append(fld_total)

	# Tab to right position
	run_tab2 = paragraph.add_run("\t")
	run_tab2.font.size = Pt(8)

	# Document ID and revision
	run_id = paragraph.add_run("[Document ID] | Rev [X]")
	set_run_font(run_id, size=8, color=RGBColor(0x66, 0x66, 0x66))


def create_table_style(doc: Document) -> None:
	"""
	Create a custom table style with borders, blue header, and alternating rows.

	python-docx has limited table style support, so we inject the style XML
	directly into the document's styles part.
	"""
	styles_element = doc.styles.element

	# Define the custom table style via raw XML
	tbl_style_xml = (
		f'<w:style {nsdecls("w")} w:type="table" w:styleId="VModelTable">'
		"  <w:name w:val=\"V-Model Table\"/>"
		"  <w:basedOn w:val=\"TableNormal\"/>"
		"  <w:uiPriority w:val=\"99\"/>"
		"  <w:pPr>"
		'    <w:spacing w:after="0" w:line="240" w:lineRule="auto"/>'
		"  </w:pPr>"
		"  <w:rPr>"
		f'    <w:rFonts w:ascii="{FONT_NAME}" w:hAnsi="{FONT_NAME}" w:eastAsia="{FONT_NAME}"/>'
		'    <w:sz w:val="20"/>'
		"  </w:rPr>"
		"  <w:tblPr>"
		'    <w:tblBorders>'
		'      <w:top w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
		'      <w:left w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
		'      <w:bottom w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
		'      <w:right w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
		'      <w:insideH w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
		'      <w:insideV w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
		'    </w:tblBorders>'
		'    <w:tblCellMar>'
		'      <w:top w:w="40" w:type="dxa"/>'
		'      <w:left w:w="80" w:type="dxa"/>'
		'      <w:bottom w:w="40" w:type="dxa"/>'
		'      <w:right w:w="80" w:type="dxa"/>'
		'    </w:tblCellMar>'
		"  </w:tblPr>"
		"  <w:tblStylePr w:type=\"firstRow\">"
		"    <w:rPr>"
		"      <w:b/>"
		'      <w:color w:val="FFFFFF"/>'
		"    </w:rPr>"
		"    <w:tcPr>"
		f'      <w:shd w:val="clear" w:color="auto" w:fill="1F4E79"/>'
		"    </w:tcPr>"
		"  </w:tblStylePr>"
		"  <w:tblStylePr w:type=\"band2Horz\">"
		"    <w:tcPr>"
		f'      <w:shd w:val="clear" w:color="auto" w:fill="{ALT_ROW_GRAY}"/>'
		"    </w:tcPr>"
		"  </w:tblStylePr>"
		"</w:style>"
	)
	styles_element.append(parse_xml(tbl_style_xml))


def generate_template() -> str:
	"""Generate the .docx template and return the output path."""
	doc = Document()

	# Page layout
	configure_page_layout(doc)

	# Styles
	configure_styles(doc)

	# Table style
	create_table_style(doc)

	# Header and footer
	build_header(doc)
	build_footer(doc)

	# Remove the default empty paragraph so the template has no body content
	for p in doc.paragraphs:
		p._element.getparent().remove(p._element)

	# Ensure output directory exists
	os.makedirs(OUTPUT_DIR, exist_ok=True)

	doc.save(OUTPUT_PATH)
	return OUTPUT_PATH


if __name__ == "__main__":
	path = generate_template()
	size_bytes = os.path.getsize(path)
	if size_bytes < 1024:
		size_str = f"{size_bytes} bytes"
	else:
		size_str = f"{size_bytes / 1024:.1f} KB"
	print(f"Template generated: {path}")
	print(f"File size: {size_str}")
	print("Done.")
